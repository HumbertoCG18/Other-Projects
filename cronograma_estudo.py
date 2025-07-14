import subprocess
import sys
import os
from datetime import datetime, timedelta
import calendar
import pandas as pd
import csv

def instalar_pacotes_necessarios():
    try:
        import docx
        print("Pacote python-docx já está instalado!")
    except ImportError:
        print("Instalando pacote python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    
    try:
        import pandas
        print("Pacote pandas já está instalado!")
    except ImportError:
        print("Instalando pacote pandas...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "pandas"])

instalar_pacotes_necessarios()

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def criar_documento():
    doc = Document()
    
    # Configurar margens da página
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo para títulos
    title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(18)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)
    
    # Estilo para subtítulos
    subtitle_style = styles.add_style('SubtitleStyle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Pt(14)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(0, 102, 153)
    
    # Estilo para cabeçalhos de nível 1
    h1_style = styles.add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 51, 102)
    
    # Estilo para cabeçalhos de nível 2
    h2_style = styles.add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 102, 153)
    
    # Estilo para cabeçalhos de nível 3
    h3_style = styles.add_style('Heading3', WD_STYLE_TYPE.PARAGRAPH)
    h3_font = h3_style.font
    h3_font.name = 'Arial'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(0, 51, 102)
    
    # Estilo para texto normal
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    
    return doc

def formatar_semana(num_semana, inicio, fim, descricao):
    if inicio.month == fim.month:
        return f"SEMANA {num_semana} ({inicio.day}-{fim.day}/{inicio.month:02d}) - {descricao}"
    else:
        return f"SEMANA {num_semana} ({inicio.day}/{inicio.month:02d}-{fim.day}/{fim.month:02d}) - {descricao}"

def gerar_cronograma_tj_rs():
    # Criar documento base
    doc = criar_documento()
    
    # Datas para o cronograma
    inicio_estudo = datetime(2025, 7, 16)
    fim_estudo = datetime(2025, 9, 30)
    periodo_estudo = (fim_estudo - inicio_estudo).days + 1
    
    # Feriados a considerar
    feriados = {
        datetime(2025, 9, 7): "Independência do Brasil",
        datetime(2025, 9, 20): "Revolução Farroupilha (RS)"
    }
    
    # Identificar dias úteis no período
    datas = []
    atual = inicio_estudo
    while atual <= fim_estudo:
        # Verificar se é dia útil (segunda a sexta)
        if atual.weekday() < 5:
            e_feriado = False
            for data_feriado in feriados:
                if atual.date() == data_feriado.date():
                    e_feriado = True
                    break
            
            if not e_feriado:
                datas.append(atual)
        
        atual += timedelta(days=1)
    
    # Criar DataFrame com as datas de estudo
    df_datas = pd.DataFrame({
        'Data': datas,
        'DiaSemana': [calendar.day_name[d.weekday()] for d in datas],
        'Mes': [calendar.month_name[d.month] for d in datas]
    })
    
    # === Título e cabeçalho ===
    title = doc.add_paragraph('CRONOGRAMA DE ESTUDOS TJ RS 2025', style='TitleStyle')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph('Técnico Judiciário - Banca FGV', style='SubtitleStyle')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Período formatado
    periodo_texto = f'Período: {inicio_estudo.strftime("%d de %B")} a {fim_estudo.strftime("%d de %B de %Y")}'
    periodo = doc.add_paragraph(periodo_texto)
    periodo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    periodo.runs[0].font.size = Pt(12)
    periodo.runs[0].italic = True
    
    doc.add_paragraph()
    
    # === Informações importantes ===
    info_heading = doc.add_heading('INFORMAÇÕES IMPORTANTES', level=1)
    info_heading.style = doc.styles['Heading1']
    
    info_list = [
        'Total de questões: 70',
        'Língua Portuguesa: 26 questões (mínimo 13 acertos) - ELIMINATÓRIO',
        'Raciocínio Lógico: 16 questões (mínimo 8 acertos) - ELIMINATÓRIO', 
        'Legislação: 16 questões - CLASSIFICATÓRIO',
        'Microinformática: 12 questões - CLASSIFICATÓRIO'
    ]
    
    for info in info_list:
        p = doc.add_paragraph('', style='List Bullet')
        runner = p.add_run(info)
        if 'ELIMINATÓRIO' in info:
            runner.font.bold = True
            runner.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    # === Distribuição de tempo ===
    doc.add_heading('DISTRIBUIÇÃO DE TEMPO POR MATÉRIA', level=1).style = doc.styles['Heading1']
    
    # Criar tabela
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    table.width = Inches(6.0)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Larguras das colunas
    widths = [Inches(3.0), Inches(1.5), Inches(1.5)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width
    
    # Cabeçalho da tabela
    hdr_cells = table.rows[0].cells
    headers = ['Matéria', '% do Tempo', 'Horas/Semana']
    
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Dados da tabela
    data = [
        ['Língua Portuguesa', '35%', '7 horas'],
        ['Raciocínio Lógico', '25%', '5 horas'],
        ['Legislação', '25%', '5 horas'],
        ['Microinformática', '15%', '3 horas']
    ]
    
    for i, row_data in enumerate(data):
        row_cells = table.rows[i+1].cells
        for j, cell_data in enumerate(row_data):
            cell = row_cells[j]
            cell.text = cell_data
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # === CRONOGRAMA DETALHADO ===
    cronograma_heading = doc.add_heading('CRONOGRAMA DETALHADO', level=1)
    cronograma_heading.style = doc.styles['Heading1']
    
    # === JULHO ===
    mes_heading = doc.add_heading('JULHO 2025 (16 a 31)', level=2)
    mes_heading.style = doc.styles['Heading2']
    
    fase_p = doc.add_paragraph('Fase: Nivelamento e Base Teórica')
    fase_p.runs[0].italic = True
    fase_p.runs[0].bold = True
    
    # Semanas de julho
    data_inicio_s1 = datetime(2025, 7, 16)
    data_fim_s1 = datetime(2025, 7, 18)
    semana1 = formatar_semana(1, data_inicio_s1, data_fim_s1, "DIAGNÓSTICO E PLANEJAMENTO")
    
    data_inicio_s2 = datetime(2025, 7, 21)
    data_fim_s2 = datetime(2025, 7, 25)
    semana2 = formatar_semana(2, data_inicio_s2, data_fim_s2, "BASE TEÓRICA I")
    
    data_inicio_s3 = datetime(2025, 7, 28)
    data_fim_s3 = datetime(2025, 7, 31)
    semana3 = formatar_semana(3, data_inicio_s3, data_fim_s3, "BASE TEÓRICA II")
    
    julho_content = f"""
{semana1}
• Quarta ({data_inicio_s1.day:02d}/{data_inicio_s1.month:02d}): Simulado diagnóstico geral + Análise de pontos fracos
• Quinta ({(data_inicio_s1 + timedelta(days=1)).day:02d}/{(data_inicio_s1 + timedelta(days=1)).month:02d}): Língua Portuguesa - Interpretação de textos (técnicas FGV)
• Sexta ({data_fim_s1.day:02d}/{data_fim_s1.month:02d}): Raciocínio Lógico - Fundamentos de lógica proposicional

{semana2}
• Segunda ({data_inicio_s2.day:02d}/{data_inicio_s2.month:02d}): Legislação - Constituição Federal (arts. 5º ao 11)
• Terça ({(data_inicio_s2 + timedelta(days=1)).day:02d}/{(data_inicio_s2 + timedelta(days=1)).month:02d}): Língua Portuguesa - Ortografia e acentuação
• Quarta ({(data_inicio_s2 + timedelta(days=2)).day:02d}/{(data_inicio_s2 + timedelta(days=2)).month:02d}): Microinformática - Windows e conceitos básicos
• Quinta ({(data_inicio_s2 + timedelta(days=3)).day:02d}/{(data_inicio_s2 + timedelta(days=3)).month:02d}): Raciocínio Lógico - Conectivos lógicos e tabelas-verdade
• Sexta ({data_fim_s2.day:02d}/{data_fim_s2.month:02d}): Legislação - Constituição Federal (arts. 37 ao 43)

{semana3}
• Segunda ({data_inicio_s3.day:02d}/{data_inicio_s3.month:02d}): Língua Portuguesa - Classes de palavras
• Terça ({(data_inicio_s3 + timedelta(days=1)).day:02d}/{(data_inicio_s3 + timedelta(days=1)).month:02d}): Raciocínio Lógico - Equivalências e negações
• Quarta ({(data_inicio_s3 + timedelta(days=2)).day:02d}/{(data_inicio_s3 + timedelta(days=2)).month:02d}): Microinformática - Pacote Office (Word e Excel)
• Quinta ({data_fim_s3.day:02d}/{data_fim_s3.month:02d}): Legislação - Código Civil (pessoas naturais e jurídicas)
"""
    
    doc.add_paragraph(julho_content)
    
    # === AGOSTO ===
    mes_heading = doc.add_heading('AGOSTO 2025', level=2)
    mes_heading.style = doc.styles['Heading2']
    
    fase_p = doc.add_paragraph('Fase: Aprofundamento e Prática')
    fase_p.runs[0].italic = True
    fase_p.runs[0].bold = True
    
    # Semanas de agosto
    data_inicio_s4 = datetime(2025, 8, 1)
    data_fim_s4 = datetime(2025, 8, 1)
    semana4 = formatar_semana(4, data_inicio_s4, data_fim_s4, "CONSOLIDAÇÃO")
    
    data_inicio_s5 = datetime(2025, 8, 4)
    data_fim_s5 = datetime(2025, 8, 8)
    semana5 = formatar_semana(5, data_inicio_s5, data_fim_s5, "APROFUNDAMENTO I")
    
    data_inicio_s6 = datetime(2025, 8, 11)
    data_fim_s6 = datetime(2025, 8, 15)
    semana6 = formatar_semana(6, data_inicio_s6, data_fim_s6, "APROFUNDAMENTO II")
    
    data_inicio_s7 = datetime(2025, 8, 18)
    data_fim_s7 = datetime(2025, 8, 22)
    semana7 = formatar_semana(7, data_inicio_s7, data_fim_s7, "QUESTÕES FGV")
    
    data_inicio_s8 = datetime(2025, 8, 25)
    data_fim_s8 = datetime(2025, 8, 29)
    semana8 = formatar_semana(8, data_inicio_s8, data_fim_s8, "LEGISLAÇÃO ESPECÍFICA")
    
    agosto_content = f"""
{semana4}
• Sexta ({data_fim_s4.day:02d}/{data_fim_s4.month:02d}): Revisão geral da semana + Exercícios

{semana5}
• Segunda ({data_inicio_s5.day:02d}/{data_inicio_s5.month:02d}): Língua Portuguesa - Concordância nominal e verbal
• Terça ({(data_inicio_s5 + timedelta(days=1)).day:02d}/{(data_inicio_s5 + timedelta(days=1)).month:02d}): Raciocínio Lógico - Argumentação lógica
• Quarta ({(data_inicio_s5 + timedelta(days=2)).day:02d}/{(data_inicio_s5 + timedelta(days=2)).month:02d}): Legislação - Código de Processo Civil (parte I)
• Quinta ({(data_inicio_s5 + timedelta(days=3)).day:02d}/{(data_inicio_s5 + timedelta(days=3)).month:02d}): Microinformática - Internet e segurança
• Sexta ({data_fim_s5.day:02d}/{data_fim_s5.month:02d}): Língua Portuguesa - Regência nominal e verbal

{semana6}
• Segunda ({data_inicio_s6.day:02d}/{data_inicio_s6.month:02d}): Raciocínio Lógico - Diagramas lógicos
• Terça ({(data_inicio_s6 + timedelta(days=1)).day:02d}/{(data_inicio_s6 + timedelta(days=1)).month:02d}): Legislação - Código de Processo Civil (parte II)
• Quarta ({(data_inicio_s6 + timedelta(days=2)).day:02d}/{(data_inicio_s6 + timedelta(days=2)).month:02d}): Língua Portuguesa - Crase
• Quinta ({(data_inicio_s6 + timedelta(days=3)).day:02d}/{(data_inicio_s6 + timedelta(days=3)).month:02d}): Microinformática - LibreOffice
• Sexta ({data_fim_s6.day:02d}/{data_fim_s6.month:02d}): Simulado parcial + Análise de desempenho

{semana7}
• Segunda ({data_inicio_s7.day:02d}/{data_inicio_s7.month:02d}): Língua Portuguesa - Resolução questões FGV
• Terça ({(data_inicio_s7 + timedelta(days=1)).day:02d}/{(data_inicio_s7 + timedelta(days=1)).month:02d}): Raciocínio Lógico - Sequências e padrões
• Quarta ({(data_inicio_s7 + timedelta(days=2)).day:02d}/{(data_inicio_s7 + timedelta(days=2)).month:02d}): Legislação - Constituição Estadual RS
• Quinta ({(data_inicio_s7 + timedelta(days=3)).day:02d}/{(data_inicio_s7 + timedelta(days=3)).month:02d}): Microinformática - Questões comentadas
• Sexta ({data_fim_s7.day:02d}/{data_fim_s7.month:02d}): Língua Portuguesa - Interpretação (questões FGV)

{semana8}
• Segunda ({data_inicio_s8.day:02d}/{data_inicio_s8.month:02d}): Legislação - Consolidação Normativa Judicial
• Terça ({(data_inicio_s8 + timedelta(days=1)).day:02d}/{(data_inicio_s8 + timedelta(days=1)).month:02d}): Raciocínio Lógico - Análise combinatória
• Quarta ({(data_inicio_s8 + timedelta(days=2)).day:02d}/{(data_inicio_s8 + timedelta(days=2)).month:02d}): Língua Portuguesa - Pontuação
• Quinta ({(data_inicio_s8 + timedelta(days=3)).day:02d}/{(data_inicio_s8 + timedelta(days=3)).month:02d}): Legislação - Estatuto dos Servidores RS
• Sexta ({data_fim_s8.day:02d}/{data_fim_s8.month:02d}): Revisão mensal + Simulado
"""
    
    doc.add_paragraph(agosto_content)
    
    # === SETEMBRO ===
    mes_heading = doc.add_heading('SETEMBRO 2025', level=2)
    mes_heading.style = doc.styles['Heading2']
    
    fase_p = doc.add_paragraph('Fase: Intensificação e Revisão Final')
    fase_p.runs[0].italic = True
    fase_p.runs[0].bold = True
    
    # Semanas de setembro
    data_inicio_s9 = datetime(2025, 9, 1)
    data_fim_s9 = datetime(2025, 9, 5)
    semana9 = formatar_semana(9, data_inicio_s9, data_fim_s9, "INTENSIVO I")
    
    data_inicio_s10 = datetime(2025, 9, 8)
    data_fim_s10 = datetime(2025, 9, 12)
    semana10 = formatar_semana(10, data_inicio_s10, data_fim_s10, "INTENSIVO II")
    
    data_inicio_s11 = datetime(2025, 9, 15)
    data_fim_s11 = datetime(2025, 9, 19)
    semana11 = formatar_semana(11, data_inicio_s11, data_fim_s11, "REVISÃO FOCADA")
    
    data_inicio_s12 = datetime(2025, 9, 22)
    data_fim_s12 = datetime(2025, 9, 26)
    semana12 = formatar_semana(12, data_inicio_s12, data_fim_s12, "RETA FINAL")
    
    data_inicio_s13 = datetime(2025, 9, 29)
    data_fim_s13 = datetime(2025, 9, 30)
    semana13 = formatar_semana(13, data_inicio_s13, data_fim_s13, "ÚLTIMOS DIAS")
    
    setembro_content = f"""
{semana9}
• Segunda ({data_inicio_s9.day:02d}/{data_inicio_s9.month:02d}): Língua Portuguesa - Redação oficial
• Terça ({(data_inicio_s9 + timedelta(days=1)).day:02d}/{(data_inicio_s9 + timedelta(days=1)).month:02d}): Raciocínio Lógico - Probabilidade básica
• Quarta ({(data_inicio_s9 + timedelta(days=2)).day:02d}/{(data_inicio_s9 + timedelta(days=2)).month:02d}): Legislação - Atos processuais e prazos
• Quinta ({(data_inicio_s9 + timedelta(days=3)).day:02d}/{(data_inicio_s9 + timedelta(days=3)).month:02d}): Microinformática - Correio eletrônico
• Sexta ({data_fim_s9.day:02d}/{data_fim_s9.month:02d}): Simulado completo (70 questões)

{semana10}
• Segunda ({data_inicio_s10.day:02d}/{data_inicio_s10.month:02d}): Língua Portuguesa - Semântica e coesão
• Terça ({(data_inicio_s10 + timedelta(days=1)).day:02d}/{(data_inicio_s10 + timedelta(days=1)).month:02d}): Raciocínio Lógico - Problemas lógicos
• Quarta ({(data_inicio_s10 + timedelta(days=2)).day:02d}/{(data_inicio_s10 + timedelta(days=2)).month:02d}): Legislação - Auxiliares da justiça
• Quinta ({(data_inicio_s10 + timedelta(days=3)).day:02d}/{(data_inicio_s10 + timedelta(days=3)).month:02d}): Revisão de erros frequentes
• Sexta ({data_fim_s10.day:02d}/{data_fim_s10.month:02d}): Língua Portuguesa - Morfossintaxe

{semana11}
• Segunda ({data_inicio_s11.day:02d}/{data_inicio_s11.month:02d}): Raciocínio Lógico - Revisão completa
• Terça ({(data_inicio_s11 + timedelta(days=1)).day:02d}/{(data_inicio_s11 + timedelta(days=1)).month:02d}): Legislação - Pontos críticos
• Quarta ({(data_inicio_s11 + timedelta(days=2)).day:02d}/{(data_inicio_s11 + timedelta(days=2)).month:02d}): Língua Portuguesa - Pegadinhas FGV
• Quinta ({(data_inicio_s11 + timedelta(days=3)).day:02d}/{(data_inicio_s11 + timedelta(days=3)).month:02d}): Microinformática - Revisão geral
• Sexta ({data_fim_s11.day:02d}/{data_fim_s11.month:02d}): Revolução Farroupilha (Feriado RS) - Estudo extra

{semana12}
• Segunda ({data_inicio_s12.day:02d}/{data_inicio_s12.month:02d}): Simulado final - Língua Portuguesa
• Terça ({(data_inicio_s12 + timedelta(days=1)).day:02d}/{(data_inicio_s12 + timedelta(days=1)).month:02d}): Simulado final - Raciocínio Lógico  
• Quarta ({(data_inicio_s12 + timedelta(days=2)).day:02d}/{(data_inicio_s12 + timedelta(days=2)).month:02d}): Simulado final - Legislação
• Quinta ({(data_inicio_s12 + timedelta(days=3)).day:02d}/{(data_inicio_s12 + timedelta(days=3)).month:02d}): Simulado final - Microinformática
• Sexta ({data_fim_s12.day:02d}/{data_fim_s12.month:02d}): Revisão de resumos e mapas mentais

{semana13}
• Segunda ({data_inicio_s13.day:02d}/{data_inicio_s13.month:02d}): Revisão de pontos críticos
• Terça ({data_fim_s13.day:02d}/{data_fim_s13.month:02d}): Revisão final e preparação psicológica
"""
    
    doc.add_paragraph(setembro_content)
    
    doc.add_page_break()
    
    # === ESTRATÉGIAS DE ESTUDO ===
    est_heading = doc.add_heading('ESTRATÉGIAS DE ESTUDO', level=1)
    est_heading.style = doc.styles['Heading1']
    
    estrategias = [
        'Sessões de 3-4 horas por noite (19h às 22h/23h)',
        'Técnica Pomodoro: 50 min estudo + 10 min descanso',
        'Fins de semana: Revisão e simulados (quando disponível)',
        'Foco em questões de provas anteriores da FGV',
        'Revisão semanal aos domingos',
        'Caderno de erros para cada matéria',
        'Mapas mentais para Legislação',
        'Flashcards para Português e Raciocínio Lógico'
    ]
    
    for est in estrategias:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        run = p.add_run(est)
        run.font.name = 'Arial'
    
    # === MATERIAIS RECOMENDADOS ===
    mat_heading = doc.add_heading('MATERIAIS RECOMENDADOS', level=1)
    mat_heading.style = doc.styles['Heading1']
    
    # Língua Portuguesa
    port_heading = doc.add_heading('Língua Portuguesa:', level=3)
    port_heading.style = doc.styles['Heading3']
    
    materiais_port = [
        'Gramática do Marcelo Rosenthal (FGV)',
        'Apostila específica FGV',
        'Questões comentadas do Qconcursos'
    ]
    for mat in materiais_port:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        run = p.add_run(mat)
    
    # Raciocínio Lógico
    rl_heading = doc.add_heading('Raciocínio Lógico:', level=3)
    rl_heading.style = doc.styles['Heading3']
    
    materiais_rl = [
        'Raciocínio Lógico FGV - Brunno Lima',
        'Questões da banca dos últimos 5 anos',
        'Curso online específico FGV'
    ]
    for mat in materiais_rl:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        run = p.add_run(mat)
    
    # Legislação
    leg_heading = doc.add_heading('Legislação:', level=3)
    leg_heading.style = doc.styles['Heading3']
    
    materiais_leg = [
        'Vade Mecum atualizado',
        'Legislação seca + questões',
        'Resumos e esquemas'
    ]
    for mat in materiais_leg:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        run = p.add_run(mat)
    
    # Microinformática
    info_heading = doc.add_heading('Microinformática:', level=3)
    info_heading.style = doc.styles['Heading3']
    
    materiais_info = [
        'Informática para Concursos - Renato da Costa',
        'Videoaulas práticas',
        'Exercícios no computador'
    ]
    for mat in materiais_info:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        run = p.add_run(mat)
    
    # === DICAS IMPORTANTES ===
    doc.add_page_break()
    dicas_heading = doc.add_heading('DICAS IMPORTANTES', level=1)
    dicas_heading.style = doc.styles['Heading1']
    
    dicas = [
        'PRIORIDADE MÁXIMA: Língua Portuguesa e Raciocínio Lógico (são eliminatórios!)',
        'Resolver no mínimo 50 questões por semana de cada matéria',
        'Fazer simulados cronometrados toda sexta-feira',
        'Manter regularidade: é melhor estudar 3h todo dia do que 10h em um dia só',
        'Cuidar da saúde: alimentação, sono e exercícios',
        'Participar de grupos de estudo online',
        'Acompanhar possíveis alterações no edital',
        'Focar nos temas mais cobrados pela FGV'
    ]
    
    for dica in dicas:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        run = p.add_run(dica)
        if 'PRIORIDADE' in dica:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    # === OBSERVAÇÕES FINAIS ===
    obs_heading = doc.add_heading('OBSERVAÇÕES FINAIS', level=1)
    obs_heading.style = doc.styles['Heading1']
    
    obs = doc.add_paragraph(
        'Este cronograma é flexível e deve ser adaptado conforme seu progresso. '
        'O importante é manter a consistência e o foco nas matérias eliminatórias. '
        'Lembre-se: a banca FGV tem características específicas, então pratique '
        'com questões anteriores dessa banca!'
    )
    obs.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Salvar documento
    doc_filename = 'Cronograma_TJ_RS_2025_Tecnico_Judiciario_v4.docx'
    doc.save(doc_filename)
    
    print(f"Cronograma criado com sucesso! Arquivo: {doc_filename}")
    
    return doc

def criar_arquivos_csv():
    # Dados das semanas
    semanas_data = [
        [1, '16-18/07', 'Diagnóstico e Planejamento', 'Simulado inicial'],
        [2, '21-25/07', 'Base Teórica I', 'Fundamentos de todas as matérias'],
        [3, '28-31/07', 'Base Teórica II', 'Consolidação dos conceitos'],
        [4, '01/08', 'Consolidação', 'Revisão e exercícios'],
        [5, '04-08/08', 'Aprofundamento I', 'Foco em temas complexos'],
        [6, '11-15/08', 'Aprofundamento II', 'Simulado parcial'],
        [7, '18-22/08', 'Questões FGV', 'Prática específica da banca'],
        [8, '25-29/08', 'Legislação Específica', 'Simulado mensal'],
        [9, '01-05/09', 'Intensivo I', 'Simulado completo'],
        [10, '08-12/09', 'Intensivo II', 'Correção de pontos fracos'],
        [11, '15-19/09', 'Revisão Focada', 'Inclui feriado estadual'],
        [12, '22-26/09', 'Reta Final', 'Simulados por matéria'],
        [13, '29-30/09', 'Últimos Ajustes', 'Preparação final']
    ]
    
    # Criar DataFrame
    df_semanas = pd.DataFrame(semanas_data, columns=['Semana', 'Período', 'Foco Principal', 'Observações'])
    
    # Salvar resumo em CSV
    resumo_filename = 'resumo_semanal_tj_rs_v4.csv'
    df_semanas.to_csv(resumo_filename, sep=';', encoding='utf-8-sig', index=False)
    
    print(f"Resumo semanal criado: {resumo_filename}")
    
    # Criar calendário de estudo
    mes_atual = 7
    ano_atual = 2025
    meses_estudo = {7: 'Julho', 8: 'Agosto', 9: 'Setembro'}
    dias_semana = ['Segunda', 'Terça', 'Quarta', 'Quinta', 'Sexta', 'Sábado', 'Domingo']
    
    calendarios = []
    
    for mes in meses_estudo.keys():
        # Obter matriz do calendário do mês
        cal = calendar.monthcalendar(ano_atual, mes)
        
        # Criar DataFrame
        df_mes = pd.DataFrame(cal)
        
        # Renomear colunas para dias da semana
        df_mes.columns = dias_semana
        
        # Adicionar informação do mês
        df_mes['Mês'] = meses_estudo[mes]
        
        calendarios.append(df_mes)
    
    # Combinar todos os meses
    df_calendario = pd.concat(calendarios)
    
    # Salvar calendário em CSV
    calendario_filename = 'calendario_estudo_tj_rs_v4.csv'
    df_calendario.to_csv(calendario_filename, sep=';', encoding='utf-8-sig', index=False)
    
    print(f"Calendário de estudo criado: {calendario_filename}")
    
    # Criar arquivo de controle de estudo diário
    datas_estudo = []
    atual = datetime(2025, 7, 16)
    fim = datetime(2025, 9, 30)
    
    while atual <= fim:
        if atual.weekday() < 5:  # Segunda a sexta
            datas_estudo.append({
                'Data': atual.strftime('%d/%m/%Y'),
                'Dia': calendar.day_name[atual.weekday()],
                'Conteúdo Planejado': '',
                'Conteúdo Estudado': '',
                'Horas Estudadas': '',
                'Questões Resolvidas': '',
                'Observações': ''
            })
        atual += timedelta(days=1)
    
    df_controle = pd.DataFrame(datas_estudo)
    
    controle_filename = 'controle_diario_estudo_tj_rs.csv'
    df_controle.to_csv(controle_filename, sep=';', encoding='utf-8-sig', index=False)
    
    print(f"Planilha de controle diário criada: {controle_filename}")

def main():
    try:
        print("Iniciando geração do cronograma de estudos para o TJ RS 2025...")
        
        # Gerar documento Word com o cronograma
        gerar_cronograma_tj_rs()
        
        # Criar arquivos CSV complementares
        criar_arquivos_csv()
        
        print("\nGeração de cronograma concluída com sucesso!")
        print("Os seguintes arquivos foram criados:")
        print("1. Cronograma_TJ_RS_2025_Tecnico_Judiciario_v4.docx (Cronograma completo)")
        print("2. resumo_semanal_tj_rs_v4.csv (Resumo por semanas)")
        print("3. calendario_estudo_tj_rs_v4.csv (Calendário mensal)")
        print("4. controle_diario_estudo_tj_rs.csv (Planilha de controle diário)")
        print("\nUtilize estes arquivos para acompanhar seu progresso de estudos.")
        
    except ImportError as e:
        print(f"ERRO: Falha ao importar bibliotecas necessárias: {e}")
        print("Execute 'pip install python-docx pandas' e tente novamente.")
    except Exception as e:
        print(f"ERRO: Ocorreu um problema durante a execução: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()